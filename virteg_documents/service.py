from docx import Document
from io import BytesIO
from django.http import HttpResponse
import os
from django.conf import settings
import datetime


class ChooseDocsVariant:
    def __init__(self, data):
        self.data = data

    def worker_method(self):
        doc_type = self.data.get("doc_type")

        if doc_type == "bill":
            return self.bill_develop()

    def bill_develop(self):
        # Путь к шаблону документа
        template_path = os.path.join(settings.BASE_DIR, 'documents/Счет_разраб_docx (1).docx')
        doc = Document(template_path)

        # Замена данных в таблице
        self.replace_main_data(doc)
        self.replace_table_data(doc)

        # Генерация ответа с готовым документом
        return self._generate_response(doc, "Счет_на_оплату.docx")

    def replace_main_data(self, doc):
        """Заменяет основные маркеры в документе."""
        replacements = {
            "{InvoiceNumber}": self.data.get("invoice_number", ""),
            "{Date}": self.data.get("date", datetime.datetime.now().strftime('%d-%m-%Y')),
            "{Customer}": self.data.get("customer", ""),
            "{INN}": self.data.get("inn", ""),
            "{KPP}": self.data.get("kpp", ""),
            "{PostalCode}": self.data.get("postal_code", ""),
            "{City}": self.data.get("city", ""),
            "{Street}": self.data.get("street", ""),
            "{Building}": self.data.get("building", ""),
            "{Office}": self.data.get("office", ""),
            "{Phone}": self.data.get("phone", ""),
            "{Total}": str(self.data.get("total", "")),
            "{AmountinRubles}": str(self.data.get("amount_in_rubles", "")),
            "{Including VAT}": str(self.data.get("including_vat", "")),
            "{TotalItems}": str(len(self.data.get("services", []))),
            "{AmountinWords}": self.data.get("amount_in_words", ""),
        }
        self.replace_text(doc, replacements)

    def replace_text(self, doc, replacements):
        """Заменяет маркеры во всех абзацах и таблицах документа."""
        # Замена в абзацах документа
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)

    def replace_table_data(self, doc):
        """Заполняет таблицу данными из запроса, заменяя маркеры на значения."""
        services = self.data.get("services", [])
        table = doc.tables[0]  # Предполагаем, что весь документ — это одна таблица
        row_offset = 24  # Начинаем со второй строки, если первая строка — заголовок

        for i, service in enumerate(services):
            table.add_row()

            row_cells = table.rows[row_offset + i].cells
            row_replacements = {
                "{Number}": service.get("number", ""),
                "{accomodation}": service.get("accomodation", ""),
                "{Quantity}": str(service.get("quantity", "")),
                "{Unit}": service.get("unit", ""),
                "{Price}": str(service.get("price", "")),
                "{Amount}": str(service.get("total", "")),
            }
            for cell in row_cells:
                self.replace_text(cell, row_replacements)

        # Замена итоговых значений после данных услуг
        self.replace_text(doc, {
            "{Total}": str(self.data.get("total", "")),
            "{AmountinRubles}": str(self.data.get("amount_in_rubles", "")),
            "{TotalItems}": str(len(services)),
            "{AmountinWords}": self.data.get("amount_in_words", ""),
        })

    def _generate_response(self, doc, filename):
        """Генерирует HTTP-ответ с файлом для скачивания."""
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename={filename}'
        return response
